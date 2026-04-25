"""
Text extraction module for PDF and DOCX files.

FIXES IN THIS VERSION:
- Added `pypdf` as primary fallback (modern replacement for deprecated PyPDF2)
- Added startup diagnostic logging so you can see WHICH libraries loaded
- Fixed silent exception swallowing — each method now logs the actual error
- Added minimum text length check per method (not just at the end)
- `extract_text_from_bytes` now logs temp file path/size for easier debugging
- All methods tried in order: pdfplumber → pypdf → PyPDF2
"""

import re
import logging
import tempfile
import os
import shutil
from typing import Optional
from pathlib import Path

# ── Library imports with explicit diagnostics ────────────────────────────────
try:
    import pdfplumber
    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    pdfplumber = None
    _PDFPLUMBER_AVAILABLE = False

try:
    from pypdf import PdfReader as _pypdf_Reader
    _PYPDF_AVAILABLE = True
except ImportError:
    _pypdf_Reader = None
    _PYPDF_AVAILABLE = False

try:
    from PyPDF2 import PdfReader as _PyPDF2_Reader
    _PYPDF2_AVAILABLE = True
except ImportError:
    _PyPDF2_Reader = None
    _PYPDF2_AVAILABLE = False

try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    Document = None
    _DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

# Log library availability once at import time — visible in Django logs
logger.info(
    f"[parser] Library availability — "
    f"pdfplumber={_PDFPLUMBER_AVAILABLE}, "
    f"pypdf={_PYPDF_AVAILABLE}, "
    f"PyPDF2={_PYPDF2_AVAILABLE}, "
    f"python-docx={_DOCX_AVAILABLE}"
)

if not _PDFPLUMBER_AVAILABLE and not _PYPDF_AVAILABLE and not _PYPDF2_AVAILABLE:
    logger.error(
        "[parser] CRITICAL: No PDF library is installed! "
        "Run: pip install pdfplumber pypdf"
    )

MAX_TEXT_LENGTH = 50_000
MIN_EXTRACTED_TEXT = 50


# ── PDF extraction ────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from PDF using all available libraries.
    Tries: pdfplumber (default) → pdfplumber (layout) → pypdf → PyPDF2
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")

    file_size = file_path.stat().st_size
    logger.info(f"[parser] Extracting PDF: {file_path.name} ({file_size} bytes)")

    # ── METHOD 1: pdfplumber default ─────────────────────────────────────────
    if _PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for idx, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        pages_text.append(page_text)
                        logger.debug(f"  pdfplumber page {idx+1}: {len(page_text)} chars")

            if pages_text:
                cleaned = clean_text("\n\n".join(pages_text))
                if len(cleaned) >= MIN_EXTRACTED_TEXT:
                    logger.info(f"[parser] ✅ pdfplumber default: {len(cleaned)} chars")
                    return cleaned[:MAX_TEXT_LENGTH]
            logger.warning("[parser] pdfplumber default returned no text, trying layout mode...")
        except Exception as e:
            logger.warning(f"[parser] pdfplumber default failed: {e}")

    # ── METHOD 2: pdfplumber with layout=True ────────────────────────────────
    if _PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    # Try layout mode first, fall back to word extraction
                    page_text = None
                    try:
                        page_text = page.extract_text(layout=True)
                    except Exception:
                        pass
                    if not page_text:
                        words = page.extract_words()
                        if words:
                            page_text = " ".join(w["text"] for w in words)
                    if page_text and page_text.strip():
                        pages_text.append(page_text)

            if pages_text:
                cleaned = clean_text("\n\n".join(pages_text))
                if len(cleaned) >= MIN_EXTRACTED_TEXT:
                    logger.info(f"[parser] ✅ pdfplumber layout: {len(cleaned)} chars")
                    return cleaned[:MAX_TEXT_LENGTH]
        except Exception as e:
            logger.warning(f"[parser] pdfplumber layout failed: {e}")

    # ── METHOD 3: pypdf (modern, actively maintained) ────────────────────────
    if _PYPDF_AVAILABLE:
        try:
            reader = _pypdf_Reader(file_path)
            pages_text = []
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages_text.append(page_text)
                    logger.debug(f"  pypdf page {idx+1}: {len(page_text)} chars")

            if pages_text:
                cleaned = clean_text("\n\n".join(pages_text))
                if len(cleaned) >= MIN_EXTRACTED_TEXT:
                    logger.info(f"[parser] ✅ pypdf: {len(cleaned)} chars")
                    return cleaned[:MAX_TEXT_LENGTH]
            logger.warning("[parser] pypdf returned no text")
        except Exception as e:
            logger.warning(f"[parser] pypdf failed: {e}")

    # ── METHOD 4: PyPDF2 (legacy fallback) ───────────────────────────────────
    if _PYPDF2_AVAILABLE:
        try:
            reader = _PyPDF2_Reader(file_path)
            pages_text = []
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages_text.append(page_text)

            if pages_text:
                cleaned = clean_text("\n\n".join(pages_text))
                if len(cleaned) >= MIN_EXTRACTED_TEXT:
                    logger.info(f"[parser] ✅ PyPDF2: {len(cleaned)} chars")
                    return cleaned[:MAX_TEXT_LENGTH]
            logger.warning("[parser] PyPDF2 returned no text")
        except Exception as e:
            logger.warning(f"[parser] PyPDF2 failed: {e}")

    # ── All methods failed ────────────────────────────────────────────────────
    installed = [
        name for name, flag in [
            ("pdfplumber", _PDFPLUMBER_AVAILABLE),
            ("pypdf", _PYPDF_AVAILABLE),
            ("PyPDF2", _PYPDF2_AVAILABLE),
        ] if flag
    ]
    if not installed:
        raise ValueError(
            "No PDF library is installed in this environment. "
            "Run: pip install pdfplumber pypdf"
        )

    raise ValueError(
        f"Could not extract text from '{file_path.name}'. "
        f"Tried: {', '.join(installed)}. "
        "The PDF may be image-based (scanned), password-protected, or corrupted. "
        "Try re-saving it from Word/Google Docs as a new PDF."
    )


# ── DOCX extraction ───────────────────────────────────────────────────────────

def extract_text_from_docx(file_path: str) -> Optional[str]:
    """Extract text from DOCX."""
    if not _DOCX_AVAILABLE:
        raise ValueError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    try:
        doc = Document(file_path)
        lines = []

        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())

        extracted_text = "\n".join(lines)
        cleaned = clean_text(extracted_text)
        logger.info(f"[parser] ✅ DOCX: {len(cleaned)} chars")
        return cleaned[:MAX_TEXT_LENGTH]

    except Exception as e:
        logger.error(f"[parser] DOCX extraction failed: {e}", exc_info=True)
        raise


# ── Bytes-based extraction (called by Django views) ──────────────────────────

def extract_text_from_bytes(file_bytes: bytes, filename: str) -> Optional[str]:
    """
    Extract text directly from file bytes.
    Called by views.py — writes bytes to a temp file, extracts, cleans up.
    """
    ext = Path(filename).suffix.lower()
    suffix = ext if ext in (".pdf", ".docx") else ".tmp"

    if suffix == ".tmp":
        raise ValueError(
            f"Unsupported file type: '{filename}'. "
            "Please upload a PDF or DOCX file."
        )

    tmp_dir = tempfile.mkdtemp(prefix="cv_extract_")
    tmp_path = os.path.join(tmp_dir, f"upload{suffix}")

    try:
        with open(tmp_path, "wb") as f:
            f.write(file_bytes)
            f.flush()
            os.fsync(f.fileno())

        actual_size = os.path.getsize(tmp_path)
        logger.info(
            f"[parser] Temp file: {tmp_path} | "
            f"received={len(file_bytes)}B, written={actual_size}B"
        )

        if actual_size == 0:
            raise ValueError("Uploaded file is empty (0 bytes).")

        if suffix == ".pdf":
            return extract_text_from_pdf(tmp_path)
        else:
            return extract_text_from_docx(tmp_path)

    finally:
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception as cleanup_err:
            logger.warning(f"[parser] Temp cleanup failed: {cleanup_err}")


# ── Text utilities ────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(lines).strip()


def extract_sections(text: str) -> dict:
    """Extract common resume sections."""
    sections = {
        "summary": "", "experience": "", "education": "",
        "skills": "", "projects": "", "certifications": "", "full_text": text,
    }

    section_keywords = {
        "summary": ["summary", "objective", "profile", "about"],
        "experience": ["experience", "employment", "work history"],
        "education": ["education", "academic", "qualifications"],
        "skills": ["skills", "technical", "competencies"],
        "projects": ["projects", "portfolio"],
        "certifications": ["certifications", "certificates", "awards"],
    }

    all_keywords = [kw for kws in section_keywords.values() for kw in kws]
    boundary = "|".join(re.escape(k) for k in all_keywords)
    boundary_pattern = rf"(?:^|\n)(?:{boundary})\s*[:\-]?\s*(?:\n|$)"

    for section_name, keywords in section_keywords.items():
        for keyword in keywords:
            start_pattern = rf"(?:^|\n){re.escape(keyword)}\s*[:\-]?\s*(?:\n|$)"
            match = re.search(start_pattern, text, re.IGNORECASE)
            if match:
                start = match.end()
                next_match = re.search(boundary_pattern, text[start:], re.IGNORECASE)
                end = start + next_match.start() if next_match else len(text)
                sections[section_name] = text[start:end].strip()
                break

    return sections


def get_text_statistics(text: str) -> dict:
    """Calculate text statistics."""
    words = text.split()
    sentences = re.split(r"[.!?]+", text)
    return {
        "word_count": len(words),
        "character_count": len(text),
        "sentence_count": len([s for s in sentences if s.strip()]),
        "average_word_length": round(len(text) / len(words), 2) if words else 0,
    }